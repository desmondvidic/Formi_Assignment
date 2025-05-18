import os
import json
from docx import Document
from pathlib import Path
from tqdm import tqdm  # Progress bar

from app.utils import chunk_text

DATA_DIR = "data"
OUTPUT_PATH = "app/kb_store.json"

def extract_text(docx_path):
    doc = Document(docx_path)
    full_text = []

    # Extract all paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)

    # Extract all tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):  # skip empty rows
                full_text.append("\t".join(row_data))

    return "\n".join(full_text)

def detect_location_from_filename(filename):
    # Assumes format like "Bangalore - Indiranagar.docx"
    parts = filename.replace(".docx", "").split("-")
    city = parts[0].strip()
    location = parts[1].strip() if len(parts) > 1 else "General"
    return city, location

def process_all_docs():
    all_chunks = []
    files = [file for file in os.listdir(DATA_DIR) if file.endswith(".docx")]

    print(f"📂 Found {len(files)} .docx files in '{DATA_DIR}'")

    for file in tqdm(files, desc="⏳ Processing files"):
        print(f"Processing file: {file}")  # <-- debug print
        path = os.path.join(DATA_DIR, file)
        city, location = detect_location_from_filename(file)
        print(f"Extracting text from: {path}")
        full_text = extract_text(path)
        print(f"Extracted text length: {len(full_text)}")

        print("Chunking text...")
        chunks = chunk_text(full_text)
        print(f"Chunks generated: {len(chunks)}")

        for chunk in chunks:
            cleaned = chunk.strip()
            if cleaned and cleaned != "\t":
                all_chunks.append({
                    "city": city,
                    "location": location,
                    "category": "general",
                    "chunk": cleaned
        })

    print(f"Writing output to {OUTPUT_PATH} ...")
    with open(OUTPUT_PATH, "w") as f:
        json.dump(all_chunks, f, indent=2)

    print(f"✅ Done! Processed {len(all_chunks)} chunks into {OUTPUT_PATH}")


if __name__ == "__main__":
    process_all_docs()
